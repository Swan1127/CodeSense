from pathlib import Path

from docx import Document

from scripts.build_guided_learning_paper_docx import build


def test_build_accepts_explicit_input_and_output(tmp_path: Path):
    source = tmp_path / "paper.md"
    target = tmp_path / "paper.docx"
    source.write_text(
        "# 测试论文\n\n## 摘要\n\n这是摘要。\n\n## 参考文献\n\n[1] 测试文献。",
        encoding="utf-8",
    )

    result = build(source, target)
    doc = Document(result)

    assert result == target
    assert target.exists()
    assert doc.paragraphs[0].text == "测试论文"
    assert doc.core_properties.author == ""
    assert doc.core_properties.last_modified_by == ""


def test_build_keeps_tables_and_relative_figures(tmp_path: Path):
    from PIL import Image

    source = tmp_path / "paper.md"
    target = tmp_path / "paper.docx"
    figure_dir = tmp_path / "figures"
    figure_dir.mkdir()
    Image.new("RGB", (1600, 900), "white").save(figure_dir / "figure.png")
    source.write_text(
        "# 测试论文\n\n"
        "![测试图](figures/figure.png)\n\n"
        "图1 测试图\n\n"
        "| 项目 | 数值 |\n|---|---:|\n| 会话 | 10 |\n",
        encoding="utf-8",
    )

    doc = Document(build(source, target))

    assert len(doc.inline_shapes) == 1
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(1, 1).text == "10"
