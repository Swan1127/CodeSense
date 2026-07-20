"""Build the Chinese guided-learning manuscript as a review-ready DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


CHINESE_BODY = "宋体"
CHINESE_HEADING = "黑体"
LATIN_FONT = "Times New Roman"
INK = RGBColor(0x00, 0x00, 0x00)
TABLE_HEADER = "F4F6F9"
CONTENT_WIDTH_DXA = 8787  # A4 width minus 3.0 cm left and 2.5 cm right margins.


def set_run_font(
    run,
    *,
    chinese: str = CHINESE_BODY,
    latin: str = LATIN_FONT,
    size: float = 12,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor = INK,
) -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), chinese)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=10.5)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = False

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_BODY)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(22)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)

    for style_name, size, before, after in (
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = LATIN_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = INK
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_HEADING)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    footer = section.footer
    footer.paragraphs[0].clear()
    add_page_number(footer.paragraphs[0])


INLINE_TOKEN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_inline_text(paragraph, text: str, *, size: float = 11, chinese=CHINESE_BODY) -> None:
    cursor = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=size, chinese=chinese)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, chinese=chinese, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, chinese=chinese, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, chinese="等线", latin="Consolas")
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, chinese=chinese)


def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(18)
    paragraph.paragraph_format.line_spacing = 1.2
    if "：" in text:
        main_title, subtitle = text.split("：", 1)
        run = paragraph.add_run(main_title)
        set_run_font(run, chinese=CHINESE_HEADING, size=18, bold=True)
        run.add_break()
        sub_run = paragraph.add_run(subtitle)
        set_run_font(sub_run, chinese=CHINESE_HEADING, size=15, bold=True)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, chinese=CHINESE_HEADING, size=18, bold=True)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.first_line_indent = Pt(0)
    add_inline_text(
        paragraph,
        text,
        size={1: 16, 2: 13, 3: 12}[level],
        chinese=CHINESE_HEADING,
    )
    for run in paragraph.runs:
        run.bold = True


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    add_inline_text(paragraph, text)


def add_equation(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    add_inline_text(paragraph, text, size=11.5, chinese="Cambria Math")


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.74)
    paragraph.paragraph_format.first_line_indent = Cm(-0.37)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208
    add_inline_text(paragraph, text)


def add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Cm(0.95)
    paragraph.paragraph_format.first_line_indent = Cm(-0.49)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208
    add_inline_text(paragraph, text)


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.74)
    paragraph.paragraph_format.right_indent = Cm(0.5)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, chinese="等线", latin="Consolas", size=10.5)
        if index < len(lines) - 1:
            run.add_break()
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F4F6")
    p_pr.append(shd)


def add_figure(doc: Document, image_path: str | Path, caption: str) -> None:
    image_path = Path(image_path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    width = 15.2
    run.add_picture(str(image_path), width=Cm(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Pt(0)
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.keep_with_next = False
    add_inline_text(cap, caption, size=10.5)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    maxima = [max(len(row[index]) for row in rows) for index in range(len(rows[0]))]
    weights = [max(5, min(value, 22)) for value in maxima]
    widths = [round(CONTENT_WIDTH_DXA * weight / sum(weights)) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    set_table_geometry(table, widths)

    table_font_size = 9.5 if len(rows[0]) <= 6 else 8.0

    for row_index, row_values in enumerate(rows):
        row = table.rows[row_index]
        if row_index == 0:
            set_repeat_table_header(row)
        for col_index, value in enumerate(row_values):
            cell = row.cells[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, TABLE_HEADER)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if col_index == 1 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            add_inline_text(paragraph, value, size=table_font_size)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True


def add_reference(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(0.74)
    paragraph.paragraph_format.first_line_indent = Cm(-0.74)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(3)
    add_inline_text(paragraph, text, size=10.5)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    add_inline_text(paragraph, text, size=10.5)


def build(manuscript_path: Path, output_path: Path) -> Path:
    lines = manuscript_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    in_code = False
    code_lines: list[str] = []
    in_references = False
    index = 0

    while index < len(lines):
        line = lines[index].rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if line.startswith("# "):
            add_title(doc, line[2:].strip())
        elif line.startswith("## "):
            heading = line[3:].strip()
            add_heading(doc, heading, 1)
            if heading == "参考文献":
                in_references = True
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
        elif re.fullmatch(r"!\[(.+)]\((.+)\)", line):
            match = re.fullmatch(r"!\[(.+)]\((.+)\)", line)
            assert match
            caption = match.group(1)
            lookahead = index + 1
            while lookahead < len(lines) and not lines[lookahead].strip():
                lookahead += 1
            if lookahead < len(lines) and re.match(r"^(图|附图)\s*[A-Za-z0-9一二三四五六七八九十]+", lines[lookahead].strip()):
                caption = lines[lookahead].strip()
                index = lookahead
            add_figure(
                doc,
                (manuscript_path.parent / match.group(2)).resolve(),
                caption,
            )
        elif line.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            parsed = [
                [cell.strip() for cell in row.strip("|").split("|")]
                for row in table_lines
            ]
            rows = [parsed[0], *parsed[2:]]
            add_markdown_table(doc, rows)
            continue
        elif line.startswith("- "):
            add_bullet(doc, line[2:].strip())
        elif re.match(r"^\d+\.\s+", line):
            add_numbered(doc, re.sub(r"^\d+\.\s+", "", line))
        elif re.match(r"^(表|附表)\s*[A-Za-z0-9一二三四五六七八九十]+", line):
            add_caption(doc, line)
        elif line.startswith("Yᵢⱼ ="):
            add_equation(doc, line)
        elif line:
            if doc.paragraphs[-1].text == "关键词":
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.space_after = Pt(8)
                add_inline_text(paragraph, line)
            elif in_references:
                add_reference(doc, line)
            else:
                add_body_paragraph(doc, line)
        index += 1

    core = doc.core_properties
    core.title = next(
        (line[2:].strip() for line in lines if line.startswith("# ")),
        manuscript_path.stem,
    )
    core.subject = "中文研究论文内部审阅稿"
    core.author = ""
    core.last_modified_by = ""
    core.comments = "内部研究稿；正式投稿前需由研究团队确认伦理程序。"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-md", required=True, type=Path)
    parser.add_argument("--output-docx", required=True, type=Path)
    args = parser.parse_args()
    print(build(args.input_md, args.output_docx))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
